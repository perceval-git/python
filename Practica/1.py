# -*- coding: utf-8 -*-

import re
from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from sys import argv
from os.path import isfile


def find_all_templates(document):
    templates = set()
    for paragraph in document.paragraphs:
        for match in re.finditer('\\{\\{.*?\\}\\}', paragraph.text):
            templates.add(match.group(0))

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                templates = templates.union(find_all_templates(cell))
                #for match in re.search('\\{\\{.*?\\}\\}', cell):
                #    templates.add(match.group(0))
    return templates


def docx_replace(document, substr, replace):
    for p in document.paragraphs:
        if substr in p.text:
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            was_replaced = False
            for i in range(len(inline)):
                if substr in inline[i].text:
                    #text = regex.sub(replace, inline[i].text)
                    text = inline[i].text.replace(substr, replace)
                    was_replaced = True
                    #print('replacing {} with {} in {}'.format(substr, replace, inline[i].text))
                    inline[i].text = text
            if not was_replaced:
                text = p.text.replace(substr, replace)
                p.text = text
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace(cell, substr , replace)


def main():
    default_filename = '/home/nicto/Документы/document.docx'
    filename = input("Enter your filename: ")
    if filename == '':
        print("The default document will be used.")
        filename = default_filename

        if len(argv) >= 2:
            filename = argv[-1]
    if not isfile(filename):
        print("No such file: {}".format(filename))
        return main()
    try:
        doc = Document(filename)
    except PackageNotFoundError:
        print("Invalid file format. You should use only \"docx\" ")
        return main()

    replacements = {}
    for template in find_all_templates(doc):
        replacement = input('Replace {} with: '.format(template))
        replacements[template] = replacement

    for template, replacement in replacements.items():
        docx_replace(doc, template, replacement)

    doc.save('/home/nicto/Документы/document-output.docx')


if __name__ == '__main__':
    while True:
        try:
            main()
        except KeyboardInterrupt:
            print("\nПродолжение работы")
